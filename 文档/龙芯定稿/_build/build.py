# -*- coding: utf-8 -*-
"""Rebuild the competition DOCX from the finalized markdown (作品报告_龙芯定稿.md).

Mirrors the markdown content faithfully INCLUDING LaTeX formulas (rendered to
images) and Python code blocks, with compact / clean tables and the original
competition template look (cover banner, header with rule, page-number footer).
"""
import os, re, sys
from PIL import Image

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import formula as F

ROOT = r"D:\Code\自动辅助驾驶仿真平台\文档\龙芯定稿"
MD = os.path.join(ROOT, "作品报告_龙芯定稿.md")
OUT = os.path.join(ROOT, "面向真车的龙芯边缘智能双冗余ADAS控制系统_作品报告_PDF四部分去重精修版.docx")
FIMG = os.path.join(ROOT, "_build", "公式img")
os.makedirs(FIMG, exist_ok=True)

CONTENT_W_CM = 15.8           # text/table content width
FONT_CN = "宋体"
FONT_CN_H = "黑体"
FONT_EN = "Times New Roman"
FONT_MONO = "Consolas"

HEADER_TEXT = "2026嵌入式大赛应用赛道作品报告"
BANNER_TEXT = "2026 嵌入式大赛 · 应用赛道作品报告"
KEYWORDS = "关键词：龙芯 2K1000LA / ADAS / 主备冗余 / MCU 安全仲裁 / CARLA-MATLAB 硬件在环"

# ----------------------------------------------------------------------------- helpers
def set_font(run, cn=FONT_CN, en=FONT_EN, size=None, bold=None, italic=None, color=None):
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en); rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn); rFonts.set(qn('w:cs'), en)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color is not None: run.font.color.rgb = RGBColor.from_string(color)

def _shd(elem, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    elem.append(shd)

def set_cell_shading(cell, fill):
    _shd(cell._tc.get_or_add_tcPr(), fill)

def set_cell_margins(cell, top=30, bottom=30, left=70, right=70):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, v in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)

def set_table_borders(table, color="9DB7C9", sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for tag in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + tag)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)

def set_table_fixed(table):
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:tblHeader'); h.set(qn('w:val'), 'true'); trPr.append(h)

def no_cell_space_after(cell):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0

def para_shading(p, fill):
    _shd(p._p.get_or_add_pPr(), fill)

def para_box(p, color="C9C9C9", sz=4, space=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for tag in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + tag)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), str(space)); e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)

# ----------------------------------------------------------------------------- inline
TOKEN = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\$[^$]+\$)')

def clean_inline_math(s):
    s = s.replace(r'\le', '≤').replace(r'\ge', '≥')
    s = s.replace(r'\times', '×').replace(r'\cdot', '·')
    s = s.replace(r'\_', '_').replace(r'\,', '').replace(r'\ ', ' ')
    s = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', s)
    s = s.replace('^2', '²').replace('^{2}', '²')
    s = s.replace('{', '').replace('}', '').replace('\\', '')
    return s

def add_inline(p, text, base_cn=FONT_CN, base_en=FONT_EN, size=None, color=None, bold_all=False):
    pos = 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()]); set_font(r, base_cn, base_en, size, bold=bold_all, color=color)
        tok = m.group(0)
        if tok.startswith('**'):
            r = p.add_run(tok[2:-2]); set_font(r, base_cn, base_en, size, bold=True, color=color)
        elif tok.startswith('`'):
            r = p.add_run(tok[1:-1]); set_font(r, base_cn, FONT_MONO, (size or 12) - 0.5, bold=bold_all, color="33475B")
        else:  # inline math
            r = p.add_run(clean_inline_math(tok[1:-1])); set_font(r, base_cn, base_en, size, italic=True, bold=bold_all, color=color)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:]); set_font(r, base_cn, base_en, size, bold=bold_all, color=color)

# ----------------------------------------------------------------------------- markdown parse
def parse_md(path):
    lines = open(path, encoding='utf-8').read().split('\n')
    blocks = []
    i = 0
    cover_done = False
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if s == '' or s == '---':
            i += 1; continue
        # cover title
        if not cover_done and s.startswith('# ') and not s.startswith('## '):
            blocks.append(('cover_title', s[2:].strip()))
            # next non-blank, non --- line = subtitle
            j = i + 1
            while j < len(lines) and lines[j].strip() in ('', '---'):
                j += 1
            if j < len(lines) and not lines[j].strip().startswith('#'):
                blocks.append(('cover_sub', lines[j].strip())); i = j + 1
            else:
                i = i + 1
            cover_done = True
            continue
        if s.startswith('#### '):
            blocks.append(('h3', s[5:].strip())); i += 1; continue
        if s.startswith('### '):
            blocks.append(('h2', s[4:].strip())); i += 1; continue
        if s.startswith('## '):
            blocks.append(('h1', s[3:].strip())); i += 1; continue
        # formula
        if s.startswith('$$'):
            buf = []
            # could be one-line $$...$$ or multi
            if s.count('$$') >= 2:
                blocks.append(('formula', s.strip('$').strip())); i += 1; continue
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('$$'):
                buf.append(lines[i]); i += 1
            i += 1
            blocks.append(('formula', '\n'.join(buf).strip())); continue
        # code fence
        if s.startswith('```'):
            i += 1; buf = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1
            blocks.append(('code', '\n'.join(buf))); continue
        # image
        mimg = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', s)
        if mimg:
            blocks.append(('image', mimg.group(2), mimg.group(1))); i += 1; continue
        # caption
        if s.startswith('〔'):
            blocks.append(('caption', s)); i += 1; continue
        # blockquote
        if s.startswith('> '):
            blocks.append(('quote', s[2:].strip())); i += 1; continue
        # table
        if s.startswith('|'):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl.append(lines[i].strip()); i += 1
            header = [c.strip() for c in tbl[0].strip('|').split('|')]
            rows = []
            for r in tbl[2:]:
                rows.append([c.strip() for c in r.strip('|').split('|')])
            blocks.append(('table', header, rows)); continue
        # bullet
        if s.startswith('- '):
            blocks.append(('bullet', s[2:].strip())); i += 1; continue
        # ordered
        mo = re.match(r'(\d+)\.\s+(.*)', s)
        if mo:
            blocks.append(('ordered', mo.group(1), mo.group(2).strip())); i += 1; continue
        # plain paragraph
        blocks.append(('para', s)); i += 1
    return blocks

# ----------------------------------------------------------------------------- column widths
def cjk_len(s):
    s = re.sub(r'\*\*|`|\$', '', s)
    return sum(2 if ord(c) > 0x2E7F else 1 for c in s)

def col_widths(header, rows, total_cm):
    n = len(header)
    maxlen = [cjk_len(header[k]) for k in range(n)]
    for r in rows:
        for k in range(min(n, len(r))):
            maxlen[k] = max(maxlen[k], cjk_len(r[k]))
    total = sum(maxlen) or n
    frac = [m / total for m in maxlen]
    lo, hi = 0.10, 0.52
    frac = [min(hi, max(lo, f)) for f in frac]
    s = sum(frac); frac = [f / s for f in frac]
    return [Cm(total_cm * f) for f in frac]

# ----------------------------------------------------------------------------- image sizing
def fig_size(path, wcap=15.0, hcap=9.3):
    w, h = Image.open(path).size
    # 竖版（整页）图：允许更高，铺满页面
    if h >= w * 0.95:
        wcap, hcap = 14.5, 21.0
    width = wcap
    height = width * h / w
    if height > hcap:
        height = hcap; width = height * w / h
    return Cm(width), Cm(height)

def formula_size(path, wcap=12.6, hcap=4.6):
    w, h = Image.open(path).size
    width = wcap
    height = width * h / w
    if height > hcap:
        height = hcap; width = height * w / h
    return Cm(width), Cm(height)

def resolve_img(p):
    p = p.strip().strip('"')
    if os.path.isabs(p):
        return p
    return os.path.join(ROOT, p.replace('/', os.sep))

# ----------------------------------------------------------------------------- builder
def build():
    blocks = parse_md(MD)
    doc = Document()

    # page setup (A4)
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.4); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.6); sec.right_margin = Cm(2.6)
    sec.header_distance = Cm(1.3); sec.footer_distance = Cm(1.3)

    # default style
    style = doc.styles['Normal']
    style.font.name = FONT_EN; style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    # header with bottom rule
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(HEADER_TEXT); set_font(r, FONT_CN, FONT_EN, 9, color="666666")
    pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4'); bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), 'BFBFBF')
    pBdr.append(bottom); hp._p.get_or_add_pPr().append(pBdr)

    # footer page number
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def fld(p, instr):
        for t, val in (('begin', None), ('instr', instr), ('end', None)):
            run = OxmlElement('w:r')
            rpr = OxmlElement('w:rPr'); rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), FONT_EN); rf.set(qn('w:hAnsi'), FONT_EN); rpr.append(rf)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18'); rpr.append(sz)
            col = OxmlElement('w:color'); col.set(qn('w:val'), '666666'); rpr.append(col)
            run.append(rpr)
            if t == 'instr':
                fc = OxmlElement('w:instrText'); fc.set(qn('xml:space'), 'preserve'); fc.text = instr; run.append(fc)
            else:
                fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), t); run.append(fc)
            p._p.append(run)
    fld(fp, ' PAGE ')

    fcount = [0]

    def new_para():
        return doc.add_paragraph()

    def body_para(text, items=None):
        p = new_para()
        pf = p.paragraph_format
        pf.line_spacing = 1.5; pf.space_after = Pt(2); pf.first_line_indent = Pt(24)
        add_inline(p, text, size=12)
        return p

    for b in blocks:
        kind = b[0]
        if kind == 'cover_title':
            for _ in range(2):
                doc.add_paragraph()
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(BANNER_TEXT); set_font(r, FONT_CN_H, FONT_EN, 15, bold=True, color="1F4E79")
            p.paragraph_format.space_after = Pt(28)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b[1]); set_font(r, FONT_CN_H, FONT_EN, 22, bold=True, color="000000")
            p.paragraph_format.space_after = Pt(14); p.paragraph_format.line_spacing = 1.3
        elif kind == 'cover_sub':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b[1]); set_font(r, "楷体", FONT_EN, 13, color="404040")
            p.paragraph_format.line_spacing = 1.4; p.paragraph_format.space_after = Pt(30)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(KEYWORDS); set_font(r, FONT_CN, FONT_EN, 11, color="333333")
            # page break before abstract
            pb = doc.add_paragraph(); run = pb.add_run(); br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); run._r.append(br)
        elif kind == 'h1':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format; pf.space_before = Pt(16); pf.space_after = Pt(10); pf.keep_with_next = True
            r = p.add_run(b[1]); set_font(r, FONT_CN_H, FONT_EN, 16, bold=True, color="1F3864")
            p.style = doc.styles['Heading 1']
            for rr in p.runs:
                set_font(rr, FONT_CN_H, FONT_EN, 16, bold=True, color="1F3864")
        elif kind == 'h2':
            p = doc.add_paragraph()
            pf = p.paragraph_format; pf.space_before = Pt(13); pf.space_after = Pt(5); pf.keep_with_next = True
            r = p.add_run(b[1]); set_font(r, FONT_CN_H, FONT_EN, 14, bold=True, color="2E5496")
            p.style = doc.styles['Heading 2']
            for rr in p.runs:
                set_font(rr, FONT_CN_H, FONT_EN, 14, bold=True, color="2E5496")
        elif kind == 'h3':
            p = doc.add_paragraph()
            pf = p.paragraph_format; pf.space_before = Pt(10); pf.space_after = Pt(4); pf.keep_with_next = True
            r = p.add_run(b[1]); set_font(r, FONT_CN_H, FONT_EN, 12.5, bold=True, color="3F6FB0")
            p.style = doc.styles['Heading 3']
            for rr in p.runs:
                set_font(rr, FONT_CN_H, FONT_EN, 12.5, bold=True, color="3F6FB0")
        elif kind == 'para':
            body_para(b[1])
        elif kind == 'bullet':
            p = new_para()
            pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_after = Pt(2)
            pf.left_indent = Cm(0.95); pf.first_line_indent = Cm(-0.5)
            rb = p.add_run('•  '); set_font(rb, FONT_CN, FONT_EN, 12, color="2E5496")
            add_inline(p, b[1], size=12)
        elif kind == 'ordered':
            p = new_para()
            pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_before = Pt(8); pf.space_after = Pt(2)
            pf.left_indent = Cm(0.0); pf.keep_with_next = True
            rn = p.add_run(b[1] + '. '); set_font(rn, FONT_CN_H, FONT_EN, 12.5, bold=True, color="1F3864")
            # render rest; bold lead handled by inline **
            add_inline(p, b[2], size=12.5, bold_all=False)
            # make the whole ordered lead visually a sub-heading: bold runs already; ensure black
        elif kind == 'quote':
            p = new_para()
            pf = p.paragraph_format; pf.left_indent = Cm(0.7); pf.line_spacing = 1.4; pf.space_after = Pt(4)
            para_shading(p, "F3F6FA")
            add_inline(p, b[1], size=11, color="444444")
        elif kind == 'code':
            code = b[1].rstrip('\n')
            p = new_para()
            pf = p.paragraph_format; pf.left_indent = Cm(0.2); pf.right_indent = Cm(0.2)
            pf.space_before = Pt(4); pf.space_after = Pt(6); pf.line_spacing = 1.12
            para_shading(p, "F4F5F7"); para_box(p, color="C9D2DC", sz=4, space=6)
            lines = code.split('\n')
            for idx, cl in enumerate(lines):
                if idx > 0:
                    br = OxmlElement('w:br'); p.add_run()._r.append(br)
                run = p.add_run(cl if cl != '' else ' ')
                set_font(run, FONT_MONO, FONT_MONO, 10.5, color="2B2B2B")
        elif kind == 'formula':
            fcount[0] += 1
            fp_img = os.path.join(FIMG, 'f%02d.png' % fcount[0])
            F.render_formula(b[1], fp_img)
            p = new_para(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
            w, h = formula_size(fp_img)
            p.add_run().add_picture(fp_img, width=w, height=h)
        elif kind == 'image':
            path = resolve_img(b[1])
            if not os.path.exists(path):
                pp = new_para(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = pp.add_run('〔缺图：%s〕' % os.path.basename(path)); set_font(r, FONT_CN, FONT_EN, 10, color="999999")
                continue
            p = new_para(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
            w, h = fig_size(path)
            p.add_run().add_picture(path, width=w, height=h)
        elif kind == 'caption':
            p = new_para(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10); p.paragraph_format.line_spacing = 1.3
            r = p.add_run(b[1]); set_font(r, FONT_CN, FONT_EN, 10.5, color="595959")
        elif kind == 'table':
            header, rows = b[1], b[2]
            n = len(header)
            widths = col_widths(header, rows, CONTENT_W_CM)
            tbl = doc.add_table(rows=1, cols=n)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            set_table_fixed(tbl); set_table_borders(tbl)
            # header
            for k, c in enumerate(header):
                cell = tbl.rows[0].cells[k]
                cell.width = widths[k]
                set_cell_shading(cell, "DCE6F2"); set_cell_margins(cell)
                pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(pp, c, size=10.5, bold_all=True, color="1F3864")
                no_cell_space_after(cell)
            repeat_header(tbl.rows[0])
            for r in rows:
                cells = tbl.add_row().cells
                for k in range(n):
                    cell = cells[k]; cell.width = widths[k]
                    set_cell_margins(cell)
                    txt = r[k] if k < len(r) else ''
                    pp = cell.paragraphs[0]
                    pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_inline(pp, txt, size=10.5)
                    no_cell_space_after(cell)
            # spacing after table
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2); sp.paragraph_format.line_spacing = 1.0
            for rr in sp.runs: pass

    # schema-clean the zoom element python-docx emits as w:val="bestFit"
    zoom = doc.settings.element.find(qn('w:zoom'))
    if zoom is not None:
        zoom.set(qn('w:percent'), '100')

    doc.save(OUT)
    print("saved:", OUT)
    print("blocks:", len(blocks), "| formulas:", fcount[0])

if __name__ == "__main__":
    build()
